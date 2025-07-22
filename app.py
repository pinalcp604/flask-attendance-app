from flask import Flask, render_template, request, redirect, url_for, session, send_file
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from email.mime.text import MIMEText
import smtplib
import os
from dotenv import load_dotenv
from io import BytesIO
import datetime

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'
EXCEL_FILE = "Attendance Test.xlsx"

load_dotenv()

# === Sheet Filtering ===


def is_valid_sheet(sheet_name):
    try:
        df = pd.read_excel(EXCEL_FILE, sheet_name=sheet_name)
        return 'Student ID' in df.columns and not sheet_name.lower().startswith('sheet')
    except:
        return False


sheet_names = [s for s in pd.ExcelFile(
    EXCEL_FILE).sheet_names if is_valid_sheet(s)]

# === Clean IDs ===


def sanitize_ids(df):
    df['Student ID'] = df['Student ID'].apply(
        lambda x: int(float(x)) if pd.notnull(x) else None)
    return df

# === LOGIN ===


@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        try:
            student_id = int(float(request.form['student_id']))
            summary = []
            selected_record = None
            selected_term = None
            week_cols = []

            for sheet in sheet_names:
                df = sanitize_ids(pd.read_excel(EXCEL_FILE, sheet_name=sheet))
                weeks = [col for col in df.columns if col.startswith("Week")]
                df[weeks] = df[weeks].clip(upper=1.0)
                df["Capped Average (%)"] = df[weeks].mean(axis=1) * 100
                df["Capped Average (%)"] = df["Capped Average (%)"].round(2)

                student_data = df[df['Student ID'] == student_id]
                if not student_data.empty:
                    record = student_data.to_dict('records')[0]
                    summary.append({"term": sheet.replace(
                        "_", " "), "average": record["Capped Average (%)"]})
                    if not selected_record:
                        selected_record = record
                        selected_term = sheet
                        week_cols = weeks

            if selected_record:
                return render_template(
                    'dashboard.html',
                    student=selected_record,
                    weeks=week_cols,
                    term=selected_term.replace("_", " "),
                    terms=sheet_names,
                    selected_term=selected_term,
                    summary=summary
                )
            return render_template('login.html', error="Student ID not found.")
        except Exception as e:
            return render_template('login.html', error=f"Error: {str(e)}")

    return render_template('login.html')

# === CHANGE TERM ===


@app.route('/term', methods=['POST'])
def change_term():
    student_id = int(float(request.form['student_id']))
    selected_term = request.form['term']
    summary = []
    selected_record = None
    week_cols = []

    for sheet in sheet_names:
        df = sanitize_ids(pd.read_excel(EXCEL_FILE, sheet_name=sheet))
        weeks = [col for col in df.columns if col.startswith("Week")]
        df[weeks] = df[weeks].clip(upper=1.0)
        df["Capped Average (%)"] = df[weeks].mean(axis=1) * 100
        df["Capped Average (%)"] = df["Capped Average (%)"].round(2)

        student_data = df[df['Student ID'] == student_id]
        if not student_data.empty:
            record = student_data.to_dict('records')[0]
            summary.append({"term": sheet.replace("_", " "),
                           "average": record["Capped Average (%)"]})
            if sheet == selected_term:
                selected_record = record
                week_cols = weeks

    if selected_record:
        return render_template('dashboard.html',
                               student=selected_record,
                               weeks=week_cols,
                               term=selected_term.replace("_", " "),
                               terms=sheet_names,
                               selected_term=selected_term,
                               summary=summary
                               )
    return render_template('login.html', error="Student ID not found.")


# === ADMIN LOGIN ===
ADMIN_USERNAME = "admin"
ADMIN_PASSWORD = "admin123"


@app.route('/admin', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        if request.form['username'] == ADMIN_USERNAME and request.form['password'] == ADMIN_PASSWORD:
            session['admin'] = True
            return redirect(url_for('admin_dashboard'))
        return render_template('admin_login.html', error="Invalid credentials")
    return render_template('admin_login.html')


@app.route('/admin/logout')
def admin_logout():
    session.pop('admin', None)
    return redirect(url_for('admin_login'))

# === ADMIN DASHBOARD ===


@app.route('/admin/dashboard', methods=['GET', 'POST'])
def admin_dashboard():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))

    selected_term = request.form.get('term') or sheet_names[0]
    student_id = request.form.get('student_id')
    student_summary = []
    student_name = None

    df = sanitize_ids(pd.read_excel(EXCEL_FILE, sheet_name=selected_term))
    week_cols = [col for col in df.columns if col.startswith("Week")]
    df[week_cols] = df[week_cols].clip(upper=1.0)
    df["Capped Average (%)"] = df[week_cols].mean(axis=1) * 100
    df["Capped Average (%)"] = df["Capped Average (%)"].round(2)
    students = df[["Student ID", "First Name", "Surname",
                   "Capped Average (%)"]].to_dict('records')

    if student_id:
        student_id = int(float(student_id))
        for sheet in sheet_names:
            df_term = sanitize_ids(pd.read_excel(EXCEL_FILE, sheet_name=sheet))
            weeks = [col for col in df_term.columns if col.startswith("Week")]
            df_term[weeks] = df_term[weeks].clip(upper=1.0)
            df_term["Capped Average (%)"] = df_term[weeks].mean(axis=1) * 100
            df_term["Capped Average (%)"] = df_term["Capped Average (%)"].round(
                2)
            student_data = df_term[df_term['Student ID'] == student_id]
            if not student_data.empty:
                record = student_data.to_dict('records')[0]
                weekly = {week: round(record[week]*100, 2) for week in weeks}
                student_summary.append({"term": sheet.replace(
                    "_", " "), "average": record["Capped Average (%)"], "weeks": weekly})
                if not student_name:
                    student_name = f"{record['First Name']} {record['Surname']}"

    return render_template(
        'admin_dashboard.html',
        students=students,
        terms=sheet_names,
        selected_term=selected_term.replace("_", " "),
        student_summary=student_summary,
        student_id=student_id,
        student_name=student_name
    )

# === WORD REPORT WITH SELECTED TERMS ===


@app.route('/admin/download_word_selected', methods=['POST'])
def download_word_selected():
    student_id = int(float(request.form['student_id']))
    selected_terms = request.form.getlist('selected_terms')
    student_summary = []
    student_name = None

    for sheet in sheet_names:
        term_name = sheet.replace("_", " ")
        if term_name not in selected_terms:
            continue

        df = sanitize_ids(pd.read_excel(EXCEL_FILE, sheet_name=sheet))
        weeks = [col for col in df.columns if col.startswith("Week")]
        df[weeks] = df[weeks].clip(upper=1.0)
        df["Capped Average (%)"] = df[weeks].mean(axis=1) * 100
        df["Capped Average (%)"] = df["Capped Average (%)"].round(2)

        student_data = df[df['Student ID'] == student_id]
        if not student_data.empty:
            record = student_data.to_dict('records')[0]
            weekly = [round(record[week]*100, 2) for week in weeks]
            student_summary.append({
                "term": term_name,
                "weeks": weekly,
                "average": record["Capped Average (%)"]
            })
            if not student_name:
                student_name = f"{record['First Name']} {record['Surname']}"

    # === Begin Word Document ===
    doc = Document()

    # Header Table
    try:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        # run.add_picture('logo1.png', width=Inches(6.2))  # Stretch across usable A4 width
        import os
        logo_path = os.path.join(os.path.dirname(
            __file__), 'static', 'images', 'logo1.PNG')

        if os.path.exists(logo_path):
            run.add_picture(logo_path, width=Inches(6.2))  # Full width banner
        else:
            run.add_text("Whitecliffe")  # Fallback text

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph("Whitecliffe").alignment = WD_ALIGN_PARAGRAPH.CENTER

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def add_horizontal_line(doc):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()

        # Create bottom border
        p_pr = p._p.get_or_add_pPr()
        border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')       # Thickness
        bottom.set(qn('w:space'), '1')    # Spacing
        bottom.set(qn('w:color'), 'auto')  # Color
        border.append(bottom)
        p_pr.append(border)

    # Call it after header
    # add_horizontal_line(doc)

    # right_cell = table.cell(0, 1)
    # p_info = right_cell.paragraphs[0]
    # _info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # un_info = p_info.add_run("67 Symonds Street\nAuckland 1010\n\n0800 800 300\nwhitecliffe.ac.nz")
    # run_info.font.size = Pt(11)
    # run_info.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%d %B %Y')}")
   # doc.add_paragraph("")
    # doc.add_paragraph(f"Student Name: {student_name}")

    # Student Name
    p_name = doc.add_paragraph()
    p_name.add_run("Student Name: ").bold = True
    p_name.add_run(student_name)

    # doc.add_paragraph(f"Student ID: {student_id}")

    # Student ID
    p_id = doc.add_paragraph()
    p_id.add_run("Student ID: ").bold = True
    p_id.add_run(str(student_id))

    # doc.add_paragraph("NSN Number: [INSERT NSN]")

    p_nsn = doc.add_paragraph()
    p_nsn.add_run("NSN Number: ").bold = True
    p_nsn.add_run("[INSERT NSN]")

    doc.add_paragraph(
        "Dear Immigration Officer,\n\nRE: Attendance Summary Letter")
    doc.add_paragraph(
        "Please be advised that the above student was enrolled in the ______________________________________ at Whitecliffe College of Arts and Design (NZQA Provider Code 8509).")
    # doc.add_paragraph("Attendance summary records:")
    p_summary = doc.add_paragraph()
    p_summary.add_run("Attendance summary records:").bold = True

    table = doc.add_table(rows=1, cols=11)
    hdr = table.rows[0].cells
    hdr[0].text = "Term"
    for i in range(1, 10):
        hdr[i].text = f"W{i}"
    hdr[10].text = "AVG"

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def set_bottom_border(cell):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')   # Line style
        bottom.set(qn('w:sz'), '6')         # Thickness
        bottom.set(qn('w:space'), '0')      # Spacing
        bottom.set(qn('w:color'), '000000')  # Black
        borders.append(bottom)
        tc_pr.append(borders)

    # Apply bottom border to each header cell
    for cell in table.rows[0].cells:
        set_bottom_border(cell)

    for term in student_summary:
        row = table.add_row().cells
        row[0].text = term["term"]
        for i in range(9):
            row[i+1].text = f"{term['weeks'][i]}%"
        row[10].text = f"{term['average']}%"

     # Set font size 9pt for all table cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph(
        "\nWe have provided this letter to confirm the student's progress and commitment towards studies.")
    doc.add_paragraph(
        "If you have any concerns relating to the above student, please do not hesitate to contact me.")

    doc.add_paragraph(
        "\nYours sincerely,\n\n"
        "Dr. Muhammad Azam (Ph.D., CITPNZ, MIITP, MPEC)\n"
        "Programme Head – Information Technology, Whitecliffe\n"
        "Level 3, Ranchhode House, Lambton Quay, Wellington\n"
        "DDI +64 44941693\n"
        "muhammada@whitecliffe.ac.nz"
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f'attendance_letter_{student_id}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/admin/generate_warning_letter/<email_type>/<int:student_id>')
def generate_warning_letter(email_type, student_id):
    # Load student data from Excel
    student_name = None
    programme = "Bachelor of Applied Information Technology (NZQF Level 7)"
    today = datetime.date.today()

    for sheet in sheet_names:
        df = sanitize_ids(pd.read_excel(EXCEL_FILE, sheet_name=sheet))
        student_data = df[df['Student ID'] == student_id]
        if not student_data.empty:
            record = student_data.to_dict('records')[0]
            student_name = f"{record['First Name']} {record['Surname']}"
            break

    if not student_name:
        return f"❌ Student ID {student_id} not found."

    template_path = f"templates/warning_template_{1 if email_type == 'warning1' else 2}.docx"
    doc = Document(template_path)

    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%d %B %Y')}")
   # doc.add_paragraph("")
    #doc.add_paragraph(f"Student Name: {student_name}")

    # Student Name
    p_name = doc.add_paragraph()
    p_name.add_run("Student Name: ").bold = True
    p_name.add_run(student_name)

    #doc.add_paragraph(f"Student ID: {student_id}")

    # Student ID
    p_id = doc.add_paragraph()
    p_id.add_run("Student ID: ").bold = True
    p_id.add_run(str(student_id))

    if email_type == "warning1":
        # === Additional Warning Content ===

        doc.add_paragraph("Re: Attendance [You are an international student currently enrolled in]", style='Normal').runs[0].bold = True

        doc.add_paragraph("")

        doc.add_paragraph(
            "You will have had conversations with faculty and/or received texts/emails about your poor attendance.\n"
            "Our records now show that your attendance at required sessions (face to face or online) has not improved since we contacted you or attempted to contact you five days ago, nor have you provided any explanation for your absence. "
            "Your success is impacted by attendance, and we would like to work with you to achieve your goals."
        )

        doc.add_paragraph("")
        doc.add_paragraph(
            "An appointment to talk about this has been made with your Programme Leader on:\n"
            "• (Date/day/time) – please reply (email/text/phone) to confirm if this will be face to face or online."
        ).runs[0].bold = True

        doc.add_paragraph("")
        doc.add_paragraph(
            "When you come, please provide the reason for your absence, along with supporting evidence.\n"
            "If you fail to attend this appointment or make contact with us, and your attendance to date falls below 90%, you may be withdrawn from your studies at Whitecliffe. "
            "Your withdrawal from the programme will affect your study visa."
        )
        doc.add_paragraph("")
        doc.add_paragraph(
            "We look forward to hearing from you as your wellbeing is very important to us, and we want to support you to complete your studies."
        )
        doc.add_paragraph("")
        doc.add_paragraph("Ngā mihi / Yours sincerely")

        doc.add_paragraph("Programme Leader")
        doc.add_paragraph("cc. Head of School")
        doc.add_paragraph("cc. Lecturer")

    if email_type == "warning2":
        doc.add_paragraph("")
        # Greeting line
        doc.add_paragraph(f"Dear or Kia Ora {student_name}")

        # Subject line
        doc.add_paragraph(
            "Re: Attendance of [Course Code and Name] in the "
            f"{programme} in which you are enrolled."
        ).runs[0].bold = True

        doc.add_paragraph("")
        doc.add_paragraph(
            "We regret to inform you that your attendance at required sessions (face to face or online) has fallen below 90%. "
            "Attendance is a condition for maintaining a student visa and Immigration New Zealand needs to be informed of any failure to attend."
        )
        doc.add_paragraph("")
        doc.add_paragraph(
            "If you do not contact Whitecliffe to discuss and, if possible, resolve your lack of attendance you will be withdrawn from the Programme, Immigration New Zealand will be informed, and you will lose your study visa. "
            "A decision will be made within the next few days by the Head of School, in consultation with the International team, who will notify you of the outcome."
        )
        doc.add_paragraph("")
        doc.add_paragraph(
            "You can gain advice from the school, or International staff."
        )
        doc.add_paragraph("")
        # Closing
        doc.add_paragraph("Ngā mihi / Yours sincerely").runs[0].bold = True

        doc.add_paragraph("Head of School").runs[0].bold = True


    

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"Warning_Letter_{email_type}_{student_id}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


if __name__ == '__main__':
    app.run(debug=True)
